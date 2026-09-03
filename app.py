import streamlit as st
import pandas as pd
import plotly.express as px
import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Configuración de la página
st.set_page_config(
    page_title="Dashboard de Gestión & Operaciones de Agencias",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Estilos CSS
st.markdown("""
    <style>
    .stApp { background-color: #F1F5F9; color: #0F172A; }
    .kpi-card { border-radius: 12px; padding: 16px 12px; text-align: center; color: white; box-shadow: 0 4px 12px rgba(15, 23, 42, 0.08); }
    .kpi-asig { background: linear-gradient(135deg, #0284C7 0%, #0369A1 100%); }
    .kpi-vis { background: linear-gradient(135deg, #16A34A 0%, #15803D 100%); }
    .kpi-novis { background: linear-gradient(135deg, #DC2626 0%, #991B1B 100%); }
    .kpi-cob { background: linear-gradient(135deg, #0D9488 0%, #115E59 100%); }
    .kpi-title { font-size: 12px; font-weight: 700; text-transform: uppercase; letter-spacing: 0.6px; opacity: 0.95; }
    .kpi-number { font-size: 34px; font-weight: 800; margin-top: 2px; }
    .section-header { background-color: #FFFFFF; padding: 12px 20px; border-radius: 10px; border-left: 5px solid #0284C7; box-shadow: 0 2px 4px rgba(0,0,0,0.04); margin-bottom: 15px; }
    h1, h2, h3 { color: #0F172A !important; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }
    [data-testid="stSidebar"] { background-color: #0F172A; }
    [data-testid="stSidebar"] * { color: #F8FAFC !important; }
    </style>
""", unsafe_allow_html=True)

ARCHIVO_EXCEL = "base_visitas.xlsx"

# -----------------------------------------------------------------------------
# LECTURA ROBUSTA DE DATOS E HISTÓRICO
# -----------------------------------------------------------------------------
def cargar_datos():
    path = ARCHIVO_EXCEL if os.path.exists(ARCHIVO_EXCEL) else ("Base_Visitas.xlsx" if os.path.exists("Base_Visitas.xlsx") else None)
    if not path:
        df = pd.DataFrame()
    else:
        df = pd.read_excel(path)
    
    if not df.empty:
        df.columns = [str(c).strip() for c in df.columns]

        if "FechaCargaMatriz" not in df.columns:
            df["FechaCargaMatriz"] = datetime.now().strftime("%Y-%m-%d")

        col_region = "Región" if "Región" in df.columns else ("Region" if "Region" in df.columns else df.columns[0])
        col_coord = "Zona" if "Zona" in df.columns else ("Coordinador" if "Coordinador" in df.columns else df.columns[1])
        col_agencia = "NombreAgencia" if "NombreAgencia" in df.columns else ("Agencia" if "Agencia" in df.columns else df.columns[3])
        col_fecha = "FechaVisita" if "FechaVisita" in df.columns else ("Fecha de Visita" if "Fecha de Visita" in df.columns else "Fecha")

        if col_fecha in df.columns:
            df["Fecha_DT"] = pd.to_datetime(df[col_fecha], errors='coerce')
            df["Visitada"] = df["Fecha_DT"].apply(lambda x: "Sí" if pd.notnull(x) else "No")
        else:
            df["Fecha_DT"] = pd.NaT
            df["Visitada"] = "No"

        cols_eval = [c for c in ["Calidad", "ProcesosTransaccionales", "Procesos Transaccionales", 
                                "ManejoEfectivoControl", "Manejo Efectivo Control", 
                                "TalentoCultura", "Talento y Cultura"] if c in df.columns]
        
        if cols_eval:
            df["PromedioGral"] = df[cols_eval].apply(pd.to_numeric, errors='coerce').mean(axis=1)
        else:
            df["PromedioGral"] = 0.0

        return df, col_agencia, col_region, col_coord, col_fecha, cols_eval
    else:
        return df, "NombreAgencia", "Región", "Zona", "FechaVisita", []

df_base, col_agencia, col_region, col_coord, col_fecha, cols_eval = cargar_datos()

if 'region_activa' not in st.session_state: st.session_state['region_activa'] = "Todas"
if 'coordinador_activo' not in st.session_state: st.session_state['coordinador_activo'] = "Todos"
if 'filtro_kpi' not in st.session_state: st.session_state['filtro_kpi'] = "Todas"

def set_filtro_kpi(valor):
    st.session_state['filtro_kpi'] = valor

# GENERACIÓN DE INFORME RESUMEN EN WORD
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generar_reporte_word(df_data, col_agencia, col_region, col_coord, col_fecha, cols_eval):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    COLOR_AZUL = RGBColor(2, 132, 199)
    COLOR_TITULO = RGBColor(15, 23, 42)
    HEX_HEADER = "0284C7"
    HEX_ALT = "F8FAFC"

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("INFORME EJECUTIVO Y RESUMEN DEL DÍA\nCONTROL DE GESTIÓN DE AGENCIAS")
    r_title.font.name = 'Segoe UI'
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_AZUL

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sub = p_sub.add_run(f"Fecha de emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    r_sub.font.size = Pt(9)
    r_sub.font.italic = True

    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Resumen Ejecutivo del Día")
    r_h1.font.name = 'Segoe UI'
    r_h1.font.color.rgb = COLOR_TITULO

    tot_asig = df_data[col_agencia].nunique() if not df_data.empty else 0
    tot_vis = df_data[df_data["Visitada"] == "Sí"][col_agencia].nunique() if not df_data.empty else 0
    tot_novis = tot_asig - tot_vis
    pct_cob = (tot_vis / tot_asig * 100) if tot_asig > 0 else 0.0

    p_body1 = doc.add_paragraph()
    p_body1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body1.add_run(
        f"El presente informe consolida el estado de la supervisión y control de gestión operativo. "
        f"A la fecha de corte, se registra un universo total de {tot_asig} agencias asignadas, de las cuales "
        f"{tot_vis} han sido supervisadas satisfactoriamente, representando un nivel de cobertura operativa del {pct_cob:.1f}%. "
        f"Asimismo, se identifican {tot_novis} agencias pendientes de visita."
    )

    table_kpi = doc.add_table(rows=2, cols=4)
    table_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_kpi = ["Agencias Asignadas", "Agencias Visitadas", "Pendientes de Visita", "% Cobertura"]
    values_kpi = [str(tot_asig), str(tot_vis), str(tot_novis), f"{pct_cob:.1f}%"]

    for i, h in enumerate(headers_kpi):
        cell = table_kpi.cell(0, i)
        set_cell_background(cell, HEX_HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    for i, v in enumerate(values_kpi):
        cell = table_kpi.cell(1, i)
        set_cell_background(cell, HEX_ALT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v)
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = COLOR_AZUL

    doc.add_paragraph()

    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Desglose Operativo por Coordinador / Zona")
    r_h2.font.name = 'Segoe UI'
    r_h2.font.color.rgb = COLOR_TITULO

    if not df_data.empty:
        df_agencias_unicas = df_data.groupby([col_coord, col_agencia], as_index=False)["Visitada"].first()
        resumen_coord = df_agencias_unicas.groupby(col_coord).agg(
            Asignadas=(col_agencia, "nunique"),
            Visitadas=("Visitada", lambda x: (x == "Sí").sum()),
            Pendientes=("Visitada", lambda x: (x == "No").sum())
        ).reset_index()

        table_coord = doc.add_table(rows=1, cols=4)
        table_coord.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = table_coord.rows[0].cells
        cols_t = ["Coordinador / Zona", "Asignadas", "Visitadas", "Pendientes"]
        for i, title in enumerate(cols_t):
            set_cell_background(hdr_cells[i], HEX_HEADER)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, row in resumen_coord.iterrows():
            row_cells = table_coord.add_row().cells
            bg = HEX_ALT if row_idx % 2 == 0 else "FFFFFF"
            for col_idx, val in enumerate([row[col_coord], row["Asignadas"], row["Visitadas"], row["Pendientes"]]):
                set_cell_background(row_cells[col_idx], bg)
                p = row_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(str(val))

    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

# -----------------------------------------------------------------------------
# PESTAÑAS PRINCIPALES
# -----------------------------------------------------------------------------
tab_dash, tab_admin = st.tabs(["📊 Dashboard de Gestión & Operaciones", "⚙️ Carga de Matriz Histórica (Admin)"])

# =============================================================================
# PESTAÑA 1: DASHBOARD
# =============================================================================
with tab_dash:
    if df_base.empty:
        st.warning("⚠️ No se ha encontrado ninguna matriz cargada. Por favor, ve a la pestaña '⚙️ Carga de Matriz Histórica (Admin)' para subir el primer archivo Excel.")
    else:
        st.sidebar.title("📌 Control de Gestión")

        fechas_cargas = sorted([str(x) for x in df_base["FechaCargaMatriz"].unique() if pd.notnull(x)], reverse=True)
        opciones_matriz = ["Última Carga (Actual)"] + fechas_cargas
        matriz_sel = st.sidebar.selectbox("📂 Seleccionar Matriz por Fecha de Carga:", opciones_matriz)

        if matriz_sel == "Última Carga (Actual)":
            df_matriz_activa = df_base[df_base["FechaCargaMatriz"] == fechas_cargas[0]] if fechas_cargas else df_base.copy()
            fecha_display = fechas_cargas[0] if fechas_cargas else "Hoy"
        else:
            df_matriz_activa = df_base[df_base["FechaCargaMatriz"] == matriz_sel]
            fecha_display = matriz_sel

        regiones = ["Todas"] + sorted([str(x) for x in df_matriz_activa[col_region].dropna().unique() if str(x).strip() != ""])
        idx_reg = regiones.index(st.session_state['region_activa']) if st.session_state['region_activa'] in regiones else 0
        region_sel = st.sidebar.selectbox("Filtrar por Región:", regiones, index=idx_reg)

        if region_sel != st.session_state['region_activa']:
            st.session_state['region_activa'] = region_sel
            st.session_state['coordinador_activo'] = "Todos"
            st.session_state['filtro_kpi'] = "Todas"
            st.rerun()

        df_region = df_matriz_activa[df_matriz_activa[col_region].astype(str) == st.session_state['region_activa']] if st.session_state['region_activa'] != "Todas" else df_matriz_activa.copy()

        coordinadores = ["Todos"] + sorted([str(x) for x in df_region[col_coord].dropna().unique() if str(x).strip() != ""])
        idx_coord = coordinadores.index(st.session_state['coordinador_activo']) if st.session_state['coordinador_activo'] in coordinadores else 0
        coord_sel = st.sidebar.selectbox("Filtrar por Coordinador/Zona:", coordinadores, index=idx_coord)

        if coord_sel != st.session_state['coordinador_activo']:
            st.session_state['coordinador_activo'] = coord_sel
            st.session_state['filtro_kpi'] = "Todas"
            st.rerun()

        if st.sidebar.button("🔄 Restablecer Filtros"):
            st.session_state['region_activa'] = "Todas"
            st.session_state['coordinador_activo'] = "Todos"
            st.session_state['filtro_kpi'] = "Todas"
            st.rerun()

        st.sidebar.markdown("---")
        st.sidebar.subheader("📄 Reporte Oficial")

        df_filtrado = df_matriz_activa.copy()
        if st.session_state['region_activa'] != "Todas":
            df_filtrado = df_filtrado[df_filtrado[col_region].astype(str) == st.session_state['region_activa']]
        if st.session_state['coordinador_activo'] != "Todos":
            df_filtrado = df_filtrado[df_filtrado[col_coord].astype(str) == st.session_state['coordinador_activo']]

        doc_bytes = generar_reporte_word(df_filtrado, col_agencia, col_region, col_coord, col_fecha, cols_eval)
        st.sidebar.download_button(
            label="📄 Descargar Informe Resumen (Word)",
            data=doc_bytes,
            file_name=f"Informe_Resumen_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        tot_asig = df_filtrado[col_agencia].nunique()
        tot_vis = df_filtrado[df_filtrado["Visitada"] == "Sí"][col_agencia].nunique()
        tot_novis = tot_asig - tot_vis
        pct_cob = (tot_vis / tot_asig * 100) if tot_asig > 0 else 0.0

        st.title("📊 Control de Gestión de Agencias")
        st.caption(f"📌 Trabajando con la Matriz Cargar del Día: **{fecha_display}** (Total filas en esta matriz: **{len(df_filtrado)}**)")

        # TARJETAS DE INDICADORES + BOTONES
        k1, k2, k3, k4 = st.columns(4)
        with k1:
            st.markdown(f'<div class="kpi-card kpi-asig"><div class="kpi-title">Agencias Asignadas</div><div class="kpi-number">{tot_asig}</div></div>', unsafe_allow_html=True)
            st.button("🔍 Filtrar Asignadas", key="btn_asig", on_click=set_filtro_kpi, args=("Todas",), use_container_width=True)

        with k2:
            st.markdown(f'<div class="kpi-card kpi-vis"><div class="kpi-title">Agencias Visitadas</div><div class="kpi-number">{tot_vis}</div></div>', unsafe_allow_html=True)
            st.button("🔍 Filtrar Visitadas", key="btn_vis", on_click=set_filtro_kpi, args=("Visitadas",), use_container_width=True)

        with k3:
            st.markdown(f'<div class="kpi-card kpi-novis"><div class="kpi-title">Agencias No Visitadas</div><div class="kpi-number">{tot_novis}</div></div>', unsafe_allow_html=True)
            st.button("🔍 Filtrar Pendientes", key="btn_novis", on_click=set_filtro_kpi, args=("No Visitadas",), use_container_width=True)

        with k4:
            st.markdown(f'<div class="kpi-card kpi-cob"><div class="kpi-title">% Cobertura</div><div class="kpi-number">{pct_cob:.1f}%</div></div>', unsafe_allow_html=True)
            st.button("🔍 Ver Todas", key="btn_cob", on_click=set_filtro_kpi, args=("Todas",), use_container_width=True)

        st.markdown("<br>", unsafe_allow_html=True)

        if st.session_state['filtro_kpi'] == "Visitadas":
            df_kpi_vista = df_filtrado[df_filtrado["Visitada"] == "Sí"]
            titulo_kpi = "📌 Detalle de Agencias Visitadas"
        elif st.session_state['filtro_kpi'] == "No Visitadas":
            df_kpi_vista = df_filtrado[df_filtrado["Visitada"] == "No"]
            titulo_kpi = "📌 Detalle de Agencias No Visitadas (Pendientes)"
        else:
            df_kpi_vista = df_filtrado.copy()
            titulo_kpi = "📌 Detalle General de Agencias Asignadas"

        st.markdown(f"### {titulo_kpi} ({df_kpi_vista[col_agencia].nunique()} agencias)")
        cols_mostrar = [c for c in [col_region, col_coord, col_agencia, "Visitada", col_fecha, "PromedioGral"] if c in df_kpi_vista.columns]
        st.dataframe(df_kpi_vista[cols_mostrar].drop_duplicates(), use_container_width=True, hide_index=True)

        st.markdown("---")

        col_izq, col_der = st.columns([1.1, 1])

        with col_izq:
            st.subheader("Resumen por Coordinador / Zona")
            df_agencias_unicas = df_filtrado.groupby([col_coord, col_agencia], as_index=False)["Visitada"].first()
            resumen_coord = df_agencias_unicas.groupby(col_coord).agg(
                AgenciasAsignadas=(col_agencia, "nunique"),
                AgenciasVisitadas=("Visitada", lambda x: (x == "Sí").sum()),
                AgenciasNoVisitadas=("Visitada", lambda x: (x == "No").sum())
            ).reset_index()

            st.dataframe(resumen_coord, use_container_width=True, hide_index=True)

        with col_der:
            v_col, nv_col = st.columns(2)
            with v_col:
                st.subheader("Agencias Visitadas")
                st.dataframe(df_filtrado[df_filtrado["Visitada"] == "Sí"][[col_region, col_agencia]].drop_duplicates(), use_container_width=True, hide_index=True, height=220)
            with nv_col:
                st.subheader("Agencias No Visitadas")
                st.dataframe(df_filtrado[df_filtrado["Visitada"] == "No"][[col_region, col_agencia]].drop_duplicates(), use_container_width=True, hide_index=True, height=220)

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-header"><h3>📈 Dashboard de Operaciones y Transacciones</h3></div>', unsafe_allow_html=True)

        if "TipoTran" in df_filtrado.columns:
            c_tr1, c_tr2, c_tr3 = st.columns(3)
            total_trans = len(df_filtrado)
            terminales_unicas = df_filtrado["Terminal"].nunique() if "Terminal" in df_filtrado.columns else 0
            tipos_unicos = df_filtrado["TipoTran"].nunique()

            c_tr1.metric("Total Transacciones Registradas", f"{total_trans:,}")
            c_tr2.metric("Terminales / Cajas Activas", f"{terminales_unicas}")
            c_tr3.metric("Tipos de Transacciones", f"{tipos_unicos}")

            g_col1, g_col2 = st.columns(2)

            with g_col1:
                st.subheader("Volumen por Tipo de Transacción")
                df_tipo = df_filtrado["TipoTran"].value_counts().reset_index()
                df_tipo.columns = ["TipoTran", "Cantidad"]
                fig_tipo = px.bar(
                    df_tipo, x="Cantidad", y="TipoTran", orientation='h',
                    color="Cantidad", color_continuous_scale="Blues", text="Cantidad"
                )
                fig_tipo.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', height=280)
                st.plotly_chart(fig_tipo, use_container_width=True)

            with g_col2:
                st.subheader("Distribución por Hora de Operación")
                if "Hora" in df_filtrado.columns:
                    df_hora = df_filtrado["Hora"].value_counts().reset_index()
                    df_hora.columns = ["Hora", "Transacciones"]
                    df_hora = df_hora.sort_values("Hora")
                    fig_hora = px.line(df_hora, x="Hora", y="Transacciones", markers=True, line_shape="spline")
                    fig_hora.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', height=280)
                    st.plotly_chart(fig_hora, use_container_width=True)

            st.subheader(f"🔍 Matriz Transaccional Completa del Día {fecha_display} ({total_trans} registros)")
            cols_trans_mostrar = [c for c in [col_region, col_coord, "Codigo", col_agencia, "Anio", "Mes", "Dia", "Semana", "Hora", "Terminal", "NombreMaq", "TipoTran"] if c in df_filtrado.columns]
            st.dataframe(df_filtrado[cols_trans_mostrar], use_container_width=True, hide_index=True)
        else:
            st.info("Cargue una matriz con registros transaccionales para ver esta sección en detalle.")

# =============================================================================
# PESTAÑA 2: MÓDULO EXCLUSIVO ADMINISTRADOR
# =============================================================================
with tab_admin:
    st.title("⚙️ Módulo Administrador: Carga e Histórico por Fecha")
    st.subheader("Gestión de Matrices Diarias")

    clave = st.text_input("🔑 Ingrese Contraseña de Administrador:", type="password")

    if clave == "admin123":
        st.success("🔓 Acceso Autorizado como Administrador")
        
        col_up1, col_up2 = st.columns([2, 1])
        
        with col_up1:
            st.markdown("### 📥 Cargar Nueva Matriz")
            fecha_subida = st.date_input("🗓️ Fecha de Carga de la Matriz:", datetime.now())
            archivo_nuevo = st.file_uploader("📂 Seleccionar Matriz Excel (.xlsx)", type=["xlsx"])

            if archivo_nuevo and st.button("💾 Guardar y Registrar Matriz de la Fecha", use_container_width=True):
                try:
                    df_nuevo = pd.read_excel(archivo_nuevo)
                    df_nuevo.columns = [str(c).strip() for c in df_nuevo.columns]

                    f_str = fecha_subida.strftime("%Y-%m-%d")
                    df_nuevo["FechaCargaMatriz"] = f_str

                    if not df_base.empty and "FechaCargaMatriz" in df_base.columns:
                        df_base_sin_fecha = df_base[df_base["FechaCargaMatriz"] != f_str]
                        df_consolidado = pd.concat([df_base_sin_fecha, df_nuevo], ignore_index=True)
                    else:
                        df_consolidado = df_nuevo

                    df_consolidado.to_excel(ARCHIVO_EXCEL, index=False)
                    
                    st.balloons()
                    st.success(f"✅ ¡Matriz del {f_str} guardada con éxito en el histórico!")
                    st.rerun()

                except Exception as e:
                    st.error(f"❌ Ocurrió un error al procesar el archivo: {e}")

        with col_up2:
            st.markdown("### 📊 Estado del Histórico")
            st.metric("Total Registros Acumulados", f"{len(df_base):,}")
            fechas_existentes = sorted([str(x) for x in df_base["FechaCargaMatriz"].unique() if pd.notnull(x)], reverse=True) if not df_base.empty and "FechaCargaMatriz" in df_base.columns else []
            st.metric("Fechas Registradas", f"{len(fechas_existentes)}")

        st.markdown("---")

        # SECCIÓN PARA ELIMINAR REGISTROS DE UNA FECHA ESPECÍFICA (DEPURACIÓN)
        st.markdown("### 🗑️ Eliminar Información de una Fecha Específica (Pruebas / Limpieza)")
        st.caption("Esta herramienta permite borrar únicamente la matriz cargada en la fecha seleccionada sin afectar a las demás fechas.")

        if fechas_existentes:
            fecha_a_eliminar = st.selectbox("Elija la fecha de carga que desea borrar:", fechas_existentes)
            
            if st.button(f"🚨 Eliminar Definitivamente Registros del {fecha_a_eliminar}"):
                df_limpio = df_base[df_base["FechaCargaMatriz"] != fecha_a_eliminar]
                df_limpio.to_excel(ARCHIVO_EXCEL, index=False)
                
                st.success(f"✅ Se han eliminado todos los registros del día {fecha_a_eliminar}.")
                st.rerun()
        else:
            st.info("No hay fechas registradas actualmente para eliminar.")

        st.markdown("---")
        st.markdown("### 🔄 Sincronización Permanente con GitHub")
        
        buffer_excel = io.BytesIO()
        with pd.ExcelWriter(buffer_excel, engine='openpyxl') as writer:
            df_base.to_excel(writer, index=False)

        st.download_button(
            label="⬇️ Descargar base_visitas.xlsx Actualizado para GitHub",
            data=buffer_excel.getvalue(),
            file_name="base_visitas.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

    elif clave != "":
        st.error("🔒 Contraseña incorrecta. Acceso restringido al Administrador.")